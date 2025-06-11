import os
import openai
import gradio as gr
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile

# טעינת מפתח OpenAI מהסביבה
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
openai.api_key = OPENAI_API_KEY

def read_transcript_from_fileobj(fileobj):
    filename = fileobj.name
    ext = os.path.splitext(filename)[-1].lower()
    if ext == '.txt':
        with open(filename, encoding="utf-8") as f:
            return f.read()
    elif ext == '.docx':
        doc = Document(filename)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)
    elif ext == '.pdf':
        text = []
        with open(filename, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n'.join(text)
    else:
        return None

def set_table_rtl(table):
    tbl_pr = table._element.xpath('w:tblPr')[0]
    bidi = OxmlElement('w:bidiVisual')
    bidi.set(qn('w:val'), '1')
    tbl_pr.append(bidi)

def set_paragraph_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph._element.set(qn('w:rtl'), '1')

def add_tasks_table_to_docx(doc, summary):
    lines = summary.split('\n')
    table_start = -1
    for i, line in enumerate(lines):
        if 'טבלת משימות' in line:
            table_start = i
            break
    if table_start == -1:
        return

    table_lines = []
    for line in lines[table_start+1:]:
        if not line.strip() or (set(line.strip()) <= set('-|')):
            continue
        table_lines.append(line)

    table_lines = [l for l in table_lines if l.strip()]
    if len(table_lines) < 2:
        return

    headers = [h.strip() for h in re.split(r'\s*\|\s*', table_lines[0]) if h.strip()]
    rows = []
    for line in table_lines[1:]:
        row = [c.strip() for c in re.split(r'\s*\|\s*', line) if c.strip()]
        if row and len(row) == len(headers):
            rows.append(row)
    if not headers or not rows:
        return

    target_order = ['שם המשימה', 'שם האחראי', 'תאריך נדרש לביצוע']
    order = []
    for col in target_order:
        if col in headers:
            order.append(headers.index(col))
        else:
            order.append(None)

    table = doc.add_table(rows=1, cols=len(target_order))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_table_rtl(table)

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(target_order):
        hdr_cells[i].text = col
    for row in rows:
        row_cells = table.add_row().cells
        for i, idx in enumerate(order):
            if idx is not None:
                row_cells[i].text = row[idx]
            else:
                row_cells[i].text = ''
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_rtl(paragraph)

def save_summary_to_word(summary, out_path):
    doc = Document()
    heading = doc.add_heading('סיכום ישיבה', 0)
    set_paragraph_rtl(heading)

    table_marker = 'טבלת משימות'
    RLM = '\u200F'  # סימן RTL
    if table_marker in summary:
        before_table = summary.split(table_marker)[0]
        for line in before_table.strip().split('\n'):
            if line.strip():
                para = doc.add_paragraph(RLM + line.strip())
                set_paragraph_rtl(para)
        add_tasks_table_to_docx(doc, summary)
    else:
        for line in summary.strip().split('\n'):
            if line.strip():
                para = doc.add_paragraph(RLM + line.strip())
                set_paragraph_rtl(para)

    doc.save(out_path)

def summarize_file(file):
    if not OPENAI_API_KEY:
        return None, "API key not found."
    transcript = read_transcript_from_fileobj(file)
    if not transcript:
        return None, "קובץ לא תקין או לא נתמך"
    prompt = (
        "סכם את הישיבה הבאה, וכתוב בנקודות עיקריות:\n"
        "1. נושאים שנדונו\n"
        "2. החלטות שהתקבלו\n"
        "3. משימות ולמי הוטלו\n"
        "4. תאריכים חשובים שהוזכרו\n\n"
        "בסיום הסיכום, צור טבלה ברורה בשם 'טבלת משימות לביצוע' הכוללת עמודות: שם המשימה, שם האחראי, תאריך נדרש לביצוע (אם לא ידוע, כתוב 'לא צוין'). "
        "אל תשאיר שורות ריקות בטבלה. אם לא היו משימות, כתוב שורה אחת בלבד: 'לא נמצאו משימות לביצוע'.\n\n"
        "תמלול הישיבה:\n"
        f"{transcript}"
    )
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "אתה עוזר חכם שמסכם ישיבות."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        summary = response.choices[0].message.content
        base, _ = os.path.splitext(os.path.basename(file.name))
        summary_filename = f"{base}_SUMMARY.docx"
        out_path = os.path.join(tempfile.gettempdir(), summary_filename)
        save_summary_to_word(summary, out_path)
        return out_path, "הסיכום מוכן להורדה"
    except Exception as e:
        return None, f"שגיאה: {e}"

# CSS RTL Custom
rtl_css = """
body, .gradio-container, label, textarea, input, .output-markdown, .output-html, .gr-description, .gradio-container .gr-description, .gradio-container .gr-description p, .gr-textbox label, .gr-box, .gr-form, .gr-interface, .gr-file label {
    direction: rtl !important;
    text-align: right !important;
    font-family: Arial, 'Noto Sans Hebrew', 'Frank Ruhl Libre', 'David', sans-serif !important;
}
"""

# לוגו ב-base64 (הדבק כאן את ה־base64 המלא)
logo_md = """
<div align="center">
  <img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAlgAAABfCAYAAAAosTT8AAANm0lEQVR4nO3dzY6kSBQAQVsNIf7/F5rsYeqcSV3Z5IN5E7PHLZoI8n/3wRfOgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAADg2bO1XdqHfG7lQ4fzRZ0o2K8eZCByXf/rF3hv3sdpwEN1bz6fK8fXPvs2Vhwl4o/GZBmb7GZIyNO34+NEZeqwpmfCrkf2cjj1+9Wz5rLmpwlmE/YHqbxG3Ug+k4v8RrHEb8xTXvxMYmcvhOrGfbv/AwWd1qgP85huquMbvTcB3mj87xl6Ps8+H9Y6B5zbsA9+zx/k41PgP2eCv2btk2F4KjcfhjPrNm4ebWJ+kxZ8zIXu7H8NzwDuW81v6eN0fKp4st3PGHcB8IexrJ9kCnStKj4o/oHzneGTiOQHyKzyHJde7Zf1ZGZnNcC94r/0nAl8QzuFur7pVQ/Hz/RM+pcHZnxZ4R/DmQw17ZkB4ZclPlMZecb5OQvX43xOoB8jP0QzWQovRG/TblZve2MzePx13xPqbbj2qzgK57fx9ztZ3b57A7/cF/UJ7V/Q0lDvSnv+QvQ/Bnpvgj/G+sj/8zGH0oF6viJtVq+x7Cj5TXFQdF/3Mzsv+6/E1M/MT9BrKuA/wz0wXwD8QzT6tckGpoP0xrkku4Z1YOYHW8g4t3fO+p74e5Dhvvt+PMdl/d8b1C6YJz5ZGeKttCejxM7z5RxOZ4ReQfxB8pv6U8T6P0u7z96CP/2Px9AGjrpOE9hZwH0A7ytHodqB6YlQHz88CKzv6M09xPXwMeQm0nQ7d1wTf5jYBOe4KveVh7EvQ70PHw6p68ZDr5zk7TeBf8SG0Bu/oScQ/El8/Fsn1bHym8ztQxHzAZqA/A5zHgO5B/Ief4ym/R7RYgU0WhHztU9i/Af8DOgLMRxyQv1zGYOivJuB8uv11OCXw/74jt5A16IfGBgf60A/Fl4rrsc3hr/kOdD1+HnMuAWQznj7j/nh5ekS9AfRngwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA8O7/AZpAdqDNLF1AAAAASUVORK5CYII=
" style="height:80px; margin-bottom: 10px;">
</div>
"""

with gr.Blocks(css=rtl_css) as demo:
    gr.Markdown(logo_md)
    gr.Markdown("<h3>מערכת הפקת סיכומי ישיבה</h3>")
    with gr.Column():
        file_input = gr.File(label="העלה קובץ תמלול (txt / docx / pdf)")
        file_output = gr.File(label="הורד את הסיכום")
        message = gr.Textbox(label="הודעה")
        submit_btn = gr.Button("סכם")

    submit_btn.click(
        summarize_file,
        inputs=file_input,
        outputs=[file_output, message]
    )

# להריץ בפורט דינמי, חובה ב-Render
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    demo.launch(server_name="0.0.0.0", server_port=port)
